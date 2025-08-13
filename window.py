from PyQt5.Qt import *
from PyQt5.uic import loadUi

import numpy as np

import pandas as pd

from docx import Document
from docx.shared import Inches

import matplotlib.pyplot as plt

from matplotlib.backends.backend_qt5agg import (NavigationToolbar2QT as NavigationToolbar)  # 导入导航栏
from matplotlib.backend_bases import Event

import calculate_model as cm

from help_window import open_help_window

import os


class CustomToolbar(NavigationToolbar):  # 自定义导航栏
    def __init__(self, canvases, parent):
        self.canvases = canvases
        self.parent = parent
        self.current_canvas = canvases[0] if canvases else None
        NavigationToolbar.__init__(self, self.current_canvas, parent)
        self.setMaximumHeight(30)

    def draw(self):
        if self.current_canvas:
            self.current_canvas.draw()

    def update(self):
        if self.current_canvas:
            self.current_canvas.update()

    def set_canvas(self, canvas):
        if canvas in self.canvases:
            self.current_canvas = canvas
            NavigationToolbar.__init__(self, self.current_canvas, self.parent, coordinates=True)
            self.canvas = canvas  # 更新内部的canvas属性
            self.update()  # 更新工具栏


class ExcelDataSelector(QDialog):
    """简化版Excel数据选择对话框，用于记录用户的选择条件"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.initUI()

    def initUI(self):
        """初始化UI界面"""
        self.setWindowTitle("数据选择")
        self.setWindowIcon(QIcon('./海洋套管评估图标.ico'))
        self.resize(470, 250)

        # 创建主布局
        main_layout = QVBoxLayout(self)

        # 创建选择项
        self.createSelectionItems(main_layout)

        # 创建按钮区域
        button_layout = QHBoxLayout()

        self.confirm_btn = QPushButton("确定")
        self.confirm_btn.clicked.connect(self.accept)
        button_layout.addWidget(self.confirm_btn)

        self.cancel_btn = QPushButton("取消")
        self.cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(self.cancel_btn)

        main_layout.addLayout(button_layout)

    def createSelectionItems(self, parent_layout):
        """创建选择项控件"""
        layout = QVBoxLayout()

        # 井的类型
        well_type_layout = QHBoxLayout()
        well_type_layout.addWidget(QLabel("井的类型:"))
        self.well_type_combo = QComboBox()
        self.well_type_combo.addItems(["气井数据", "油井数据"])
        well_type_layout.addWidget(self.well_type_combo)
        layout.addLayout(well_type_layout)

        # 磨损面积模型
        wear_model_layout = QHBoxLayout()
        wear_model_layout.addWidget(QLabel("磨损面积模型:"))
        self.wear_model_combo = QComboBox()
        self.wear_model_combo.addItems(["套管磨损面积(线性关系)"])
        wear_model_layout.addWidget(self.wear_model_combo)
        layout.addLayout(wear_model_layout)

        # 环境类型
        environment_layout = QHBoxLayout()
        environment_layout.addWidget(QLabel("环境类型:"))
        self.environment_combo = QComboBox()
        self.environment_combo.addItems(["硫化氢或二氧化碳腐蚀环境"])
        environment_layout.addWidget(self.environment_combo)
        layout.addLayout(environment_layout)

        # 地层类型
        formation_layout = QHBoxLayout()
        formation_layout.addWidget(QLabel("地层类型:"))
        self.formation_combo = QComboBox()
        self.formation_combo.addItems(["非塑性蠕变地层", "塑性蠕变地层"])
        formation_layout.addWidget(self.formation_combo)
        layout.addLayout(formation_layout)

        parent_layout.addLayout(layout)

    def getSelections(self):
        """获取用户选择的结果"""
        return {
            "well_type": self.well_type_combo.currentText(),
            "wear_model": self.wear_model_combo.currentText(),
            "environment": self.environment_combo.currentText(),
            "formation": self.formation_combo.currentText()
        }


def fill_template_with_results(template_path, result_dict, chart_path, output_path):
    """
    使用计算结果填充Word模板

    参数:
    template_path: 模板文件路径
    result_dict: 计算结果字典
    chart_path: 图表图片路径
    output_path: 输出文件路径
    """
    try:
        # 打开模板文档
        doc = Document(template_path)

        # 1. 填充文本结果
        for paragraph in doc.paragraphs:
            # 替换冲蚀速率
            if "##ABLATION_VELOCITY##" in paragraph.text:
                paragraph.text = paragraph.text.replace("##ABLATION_VELOCITY##", f"{result_dict['冲蚀速率']:.3f}")

            # 替换长期腐蚀速率
            if "##CORROSION_VELOCITY##" in paragraph.text:
                paragraph.text = paragraph.text.replace("##CORROSION_VELOCITY##", f"{result_dict['长期腐蚀速率']:.4f}")

            # 替换最大外压力
            if "##MAX_EXPRESSURE##" in paragraph.text:
                paragraph.text = paragraph.text.replace("##MAX_EXPRESSURE##", f"{result_dict['最大外压力']:.2f}")

            # 替换最大内压力
            if "##MAX_INPRESSURE##" in paragraph.text:
                paragraph.text = paragraph.text.replace("##MAX_INPRESSURE##", f"{result_dict['最大内压力']:.2f}")

            # 替换剩余抗外挤强度
            if "##RESI_EXRESISTANCE_EXTRUSION##" in paragraph.text:
                paragraph.text = paragraph.text.replace("##RESI_EXRESISTANCE_EXTRUSION##", f"{result_dict['剩余抗外挤强度']:.2f}")

            # 替换剩余抗内挤强度
            if "##RESI_INRESISTANCE_EXTRUSION##" in paragraph.text:
                paragraph.text = paragraph.text.replace("##RESI_INRESISTANCE_EXTRUSION##", f"{result_dict['剩余抗内挤强度']:.2f}")

            # 替换安全等级
            if "##SAFT_LEVEL##" in paragraph.text:
                paragraph.text = paragraph.text.replace("##SAFT_LEVEL##", f"{result_dict['安全等级']}")

        # 处理表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # 替换冲蚀速率
                        if "##ABLATION_VELOCITY##" in paragraph.text:
                            paragraph.text = paragraph.text.replace("##ABLATION_VELOCITY##",
                                                                    f"{result_dict['冲蚀速率']:.3f}")

                        # 替换长期腐蚀速率
                        if "##CORROSION_VELOCITY##" in paragraph.text:
                            paragraph.text = paragraph.text.replace("##CORROSION_VELOCITY##",
                                                                    f"{result_dict['长期腐蚀速率']:.4f}")

                        # 替换最大外压力
                        if "##MAX_EXPRESSURE##" in paragraph.text:
                            paragraph.text = paragraph.text.replace("##MAX_EXPRESSURE##",
                                                                    f"{result_dict['最大外压力']:.2f}")

                        # 替换最大内压力
                        if "##MAX_INPRESSURE##" in paragraph.text:
                            paragraph.text = paragraph.text.replace("##MAX_INPRESSURE##",
                                                                    f"{result_dict['最大内压力']:.2f}")

                        # 替换剩余抗外挤强度
                        if "##RESI_EXRESISTANCE_EXTRUSION##" in paragraph.text:
                            paragraph.text = paragraph.text.replace("##RESI_EXRESISTANCE_EXTRUSION##",
                                                                    f"{result_dict['剩余抗外挤强度']:.2f}")

                        # 替换剩余抗内挤强度
                        if "##RESI_INRESISTANCE_EXTRUSION##" in paragraph.text:
                            paragraph.text = paragraph.text.replace("##RESI_INRESISTANCE_EXTRUSION##",
                                                                    f"{result_dict['剩余抗内挤强度']:.2f}")

                        # 替换安全等级
                        if "##SAFT_LEVEL##" in paragraph.text:
                            paragraph.text = paragraph.text.replace("##SAFT_LEVEL##", f"{result_dict['安全等级']}")

        # 遍历5个图表占位符
        for i in range(1, 6):
            # 定义当前图表的占位符文本和对应的图片路径
            target_text = f"##CHART_{i}##"

            # 标记是否找到占位符
            found = False

            # 遍历文档段落查找占位符
            for paragraph in doc.paragraphs:
                if target_text in paragraph.text:
                    # 清空占位符文本
                    paragraph.text = paragraph.text.replace(target_text, "")
                    # 插入对应图片
                    run = paragraph.add_run()
                    run.add_picture(chart_path[i - 1], width=Inches(3.5))
                    found = True
                    break  # 找到后退出段落循环

            # 如果没找到当前占位符，可根据需求处理
            if not found:
                # 示例：在文档末尾添加图片
                new_paragraph = doc.add_paragraph()
                new_paragraph.add_run().add_picture(chart_path[i - 1], width=Inches(6))
                # 可选：添加提示文本
                new_paragraph.add_run(f"（自动添加的图表 {i}，未找到对应占位符）")

        # 保存填充后的文档
        doc.save(output_path)
        return True, "文档填充成功", output_path

    except Exception as e:
        return False, f"错误：{str(e)}", None


class MainWindow(QMainWindow):
    def __init__(self):
        super(MainWindow, self).__init__()

        # loadUi("D:\Python3\Qt_Ui_designer\main_window1123.ui", self)
        # 获取当前脚本的目录
        current_dir = os.path.dirname(os.path.abspath(__file__))
        # 拼接相对路径
        ui_file_path = os.path.join(current_dir, "main_window1123.ui")
        # 加载 UI 文件
        loadUi(ui_file_path, self)

        # 获取所有 MplWidget 实例
        self.mpl_widgets = [
            self.widget.canvas,
            self.widget_2.canvas,
            self.widget_5.canvas,
            self.widget_6.canvas,
            self.widget_7.canvas,
            self.widget_8.canvas,
            self.widget_9.canvas,
            self.graphicsView_3.canvas
        ]

        self.toolbar = CustomToolbar(self.mpl_widgets, self)  # 添加导航栏
        self.addToolBar(self.toolbar)

        self.centralwidget.setContentsMargins(11, 11, 11, 11)  # 设置窗口边距

        self.pushButton.clicked.connect(self.gas_well_ablation_graph)
        self.pushButton_2.clicked.connect(self.oil_well_ablation_graph)
        self.pushButton_3.clicked.connect(self.wear_model_line_graph)
        self.pushButton_5.clicked.connect(self.corrode_model_graph)
        self.pushButton_7.clicked.connect(self.noplasticity_effective_external_squeeze_pressure_graph)
        self.pushButton_8.clicked.connect(self.plasticity_effective_external_squeeze_pressure_graph)
        self.pushButton_9.clicked.connect(self.gas_effective_internal_pressure_graph)
        self.pushButton_10.clicked.connect(self.oil_effective_internal_squeeze_pressure_graph)
        self.pushButton_11.clicked.connect(self.effective_internal_pressure_failure_condition_calculate)

        self.actionopen.triggered.connect(self.open_excel_file)
        self.actionsave.triggered.connect(self.export_to_template)

        self.menu_2.aboutToShow.connect(open_help_window)

        # 设定字体为微软雅黑
        plt.rcParams['font.sans-serif'] = ['Microsoft Yahei']

        self.widget.canvas.axes.set_xlabel('套管横截面积（m2）')  # 设置图像标签
        self.widget.canvas.axes.set_ylabel('侵蚀深度（mm/year）')
        self.widget.canvas.axes.set_title(' Tulsa angle dependent model')

        self.widget_2.canvas.axes.set_xlabel('钻速（m/h）')  # 设置图像标签
        self.widget_2.canvas.axes.set_ylabel('套管磨损面积（m^2）')
        self.widget_2.canvas.axes.set_title(' 磨损模型 ')

        self.widget_5.canvas.axes.set_xlabel('套管横截面积（m2）')  # 设置图像标签
        self.widget_5.canvas.axes.set_ylabel('壁厚（mm）')
        self.widget_5.canvas.axes.set_title(' DNV模型')

        self.widget_6.canvas.axes.set_xlabel('计算点深度（m）')  # 设置图像标签
        self.widget_6.canvas.axes.set_ylabel('有效外挤压力（Mpa）')
        self.widget_6.canvas.axes.set_title(' 有效外挤压力模型 ')

        self.widget_7.canvas.axes.set_xlabel('计算点深度（m）')  # 设置图像标签
        self.widget_7.canvas.axes.set_ylabel('有效外挤压力（MPa）')
        self.widget_7.canvas.axes.set_title(' 有效外挤压力模型 ')

        self.widget_8.canvas.axes.set_xlabel('计算点深度（m）')  # 设置图像标签
        self.widget_8.canvas.axes.set_ylabel('有效内压力（MPa）')
        self.widget_8.canvas.axes.set_title(' 有效内挤压力模型 ')

        self.widget_9.canvas.axes.set_xlabel('计算点深度（m）')  # 设置图像标签
        self.widget_9.canvas.axes.set_ylabel('有效内压力（MPa）')
        self.widget_9.canvas.axes.set_title(' 有效内挤压力模型 ')

        self.graphicsView_3.canvas.axes.set_xlabel('温度')  # 设置图像标签
        self.graphicsView_3.canvas.axes.set_ylabel('腐蚀速率')
        self.graphicsView_3.canvas.axes.set_title('腐蚀预测模型')

        self.d_c = 0  # 冲蚀损失壁厚
        self.d_m = 0  # 磨损损失壁厚
        self.d_f = 0  # 腐蚀损失壁厚
        self.P_ce = 0  # 最大外压力
        self.P_bh = 0  # 最大内压力

        self.temp_image_path = ["chart_1.png", "chart_2.png", "chart_3.png", "chart_4.png", "chart_5.png"]
        self.word_path = "计算结果报告.docx"
        self.calc_result = {
            "冲蚀速率": None, "长期腐蚀速率": None, "最大外压力": None,
            "最大内压力": None, "剩余抗外挤强度": None, "剩余抗内挤强度": None,
            "安全等级": None
        }
        self.last_template_output = None  # 记录最后一次模板导出的文件路径

        # 连接标签页切换信号
        self.tabWidget.currentChanged.connect(self.on_tab_changed)
        self.tabWidget_2.currentChanged.connect(self.on_tab_changed)
        self.toolBox.currentChanged.connect(self.on_tab_changed)
        self.toolBox_4.currentChanged.connect(self.on_tab_changed)
        self.toolBox_5.currentChanged.connect(self.on_tab_changed)

    def on_tab_changed(self):
        """标签页切换时更新工具栏画布（增加调试信息）"""
        try:
            if self.tabWidget.currentIndex() == 0 and self.toolBox.currentIndex() == 0:
                self.toolbar.set_canvas(self.mpl_widgets[0])
            elif self.tabWidget.currentIndex() == 0 and self.toolBox.currentIndex() == 1:
                self.toolbar.set_canvas(self.mpl_widgets[3])
            elif self.tabWidget.currentIndex() == 1:
                self.toolbar.set_canvas(self.mpl_widgets[1])
            elif self.tabWidget.currentIndex() == 2:
                self.toolbar.set_canvas(self.mpl_widgets[6])
            elif self.tabWidget.currentIndex() == 3 and self.toolBox_4.currentIndex() == 0 and self.tabWidget_2.currentIndex() == 0:
                self.toolbar.set_canvas(self.mpl_widgets[3])
            elif self.tabWidget.currentIndex() == 3 and self.toolBox_4.currentIndex() == 1 and self.tabWidget_2.currentIndex() == 0:
                self.toolbar.set_canvas(self.mpl_widgets[4])
            elif self.tabWidget.currentIndex() == 3 and self.toolBox_5.currentIndex() == 0 and self.tabWidget_2.currentIndex() == 1:
                self.toolbar.set_canvas(self.mpl_widgets[5])
            elif self.tabWidget.currentIndex() == 3 and self.toolBox_5.currentIndex() == 1 and self.tabWidget_2.currentIndex() == 1:
                self.toolbar.set_canvas(self.mpl_widgets[6])
            else:
                self.toolbar.set_canvas(self.mpl_widgets[0])
        except Exception as e:
            import logging
            logging.error(f"Error occurred in on_tab_changed: {e}")
            # 可以选择恢复默认画布或进行其他处理
            # self.toolbar.set_canvas(self.mpl_widgets[0])

    def gas_well_ablation_graph(self):  # 更新图像数据
        """
        更新图形界面中的数据。

        """
        # 气井 冲蚀模型 计算
        if self.comboBox.currentText() == '圆形':
            F_s = 0.2
        elif self.comboBox.currentText() == '半圆形':
            F_s = 0.53
        elif self.comboBox.currentText() == '角形':
            F_s = 1

        gas_well_ablation = cm.Gas_well_ablation_model(
            m=self.doubleSpinBox_2.value(),
            F=F_s,
            a=self.doubleSpinBox_5.value(),
            p=self.doubleSpinBox_67.value(),
            s=self.doubleSpinBox_68.value(),
            B=220,
            t=self.doubleSpinBox_69.value(),
            v=self.doubleSpinBox_3.value()
        )

        # 计算冲蚀壁厚
        self.d_c = gas_well_ablation.d_c
        print(f"d_c = {self.d_c} mm\n")

        length_of_signal = 50
        t = np.linspace(0.01, 0.25, length_of_signal)
        FS_signal = []
        for i in range(length_of_signal):
            FS = cm.Gas_well_ablation_model(
                m=self.doubleSpinBox_2.value(),
                F=F_s,
                a=self.doubleSpinBox_5.value(),
                p=self.doubleSpinBox_67.value(),
                s=t[i],
                B=220,
                t=self.doubleSpinBox_69.value(),
                v=self.doubleSpinBox_3.value()
            )
            FS_signal.append(FS.E)

        print(FS_signal)

        # 设定字体为微软雅黑
        plt.rcParams['font.sans-serif'] = ['Microsoft Yahei']

        # 清除当前图形以便绘制新的信号
        self.widget.canvas.axes.clear()
        # 绘制信号
        self.widget.canvas.axes.plot(t, FS_signal)
        # 设置x轴标签
        self.widget.canvas.axes.set_xlabel('管道截面面积（m2）')  # 设置图像标签
        # 设置y轴标签
        self.widget.canvas.axes.set_ylabel('侵蚀速率（mm/year）')
        # 设置图形标题
        self.widget.canvas.axes.set_title(' Tulsa angle dependent model')

        # 绘制图形
        self.widget.canvas.draw()

        # 获取Figure对象
        fig = self.widget.canvas.figure
        # 保存为图片文件
        fig.savefig(self.temp_image_path[0], dpi=300, bbox_inches='tight')

        # 显示冲蚀速率
        self.lineEdit_9.setText(f"{gas_well_ablation.E * 10 ** 4:.3f}")

        self.calc_result['冲蚀速率'] = gas_well_ablation.E * 10 ** 4

    def oil_well_ablation_graph(self):  # 更新图像数据
        """
        更新图形界面中的数据。

        """
        # 油井 冲蚀模型 计算
        oil_well_ablation = cm.Oil_well_ablation_model(
            m_p=self.doubleSpinBox_8.value(),  # 砂的流量
            U_p=self.doubleSpinBox_20.value(),  # 粒子撞击速度
            rho_t=self.doubleSpinBox_19.value(),  # 目标材料密度
            A_pipe=self.doubleSpinBox_17.value(),  # 管道的横截面积
            alpha=self.doubleSpinBox_13.value(),  # 冲蚀角度
            rho_m=self.doubleSpinBox_21.value(),  # 液体混合物密度
            dp=self.doubleSpinBox_64.value(),  # 颗粒直径
            t_c=self.doubleSpinBox_12.value()  # 冲蚀时间
        )

        # 计算冲蚀壁厚
        self.d_c = oil_well_ablation.d_c
        print(f"d_c = {self.d_c} mm\n")

        length_of_signal = 50
        t = np.linspace(0.01, 0.25, length_of_signal)
        FS_signal = []
        for i in range(length_of_signal):
            FS = cm.Oil_well_ablation_model(
                m_p=self.doubleSpinBox_8.value(),  # 砂的流量
                U_p=self.doubleSpinBox_20.value(),  # 粒子撞击速度
                rho_t=self.doubleSpinBox_19.value(),  # 目标材料密度
                A_pipe=t[i],  # 管道的横截面积
                alpha=self.doubleSpinBox_13.value(),  # 冲蚀角度
                rho_m=self.doubleSpinBox_21.value(),  # 液体混合物密度
                dp=self.doubleSpinBox_64.value(),  # 颗粒直径
                t_c=self.doubleSpinBox_12.value()  # 冲蚀时间
            )
            FS_signal.append(FS.E_cl)

        print(FS_signal)

        # 设定字体为微软雅黑
        plt.rcParams['font.sans-serif'] = ['Microsoft Yahei']

        # 清除当前图形以便绘制新的信号
        self.widget_5.canvas.axes.clear()
        # 绘制信号
        self.widget_5.canvas.axes.plot(t, FS_signal)
        # 设置x轴标签
        self.widget_5.canvas.axes.set_xlabel('管道截面面积（m2）')  # 设置图像标签
        # 设置y轴标签
        self.widget_5.canvas.axes.set_ylabel('侵蚀速率（mm/year）')
        # 设置图形标题
        self.widget_5.canvas.axes.set_title(' 弯管冲蚀模型 ')

        # 绘制图形
        self.widget_5.canvas.draw()

        # 保存为图片文件
        self.widget_5.canvas.figure.savefig(self.temp_image_path[0], dpi=300, bbox_inches='tight')

        # 显示冲蚀速率
        self.lineEdit_10.setText(f"{oil_well_ablation.E_cl * 10 ** 4:.3f}")

        self.calc_result['冲蚀速率'] = oil_well_ablation.E_cl * 10 ** 4

    def wear_model_line_graph(self):  # 更新图像数据
        """
        更新图形界面中的数据。

        """
        # 磨损模型
        wear_model_line = cm.Wear_model_line(
            mu=self.doubleSpinBox_26.value(),
            n=self.doubleSpinBox_27.value(),
            f_w=self.doubleSpinBox_25.value(),
            D=self.doubleSpinBox_24.value(),
            L_m=self.doubleSpinBox_22.value(),
            v_rop=self.doubleSpinBox_23.value(),
            Rc=self.doubleSpinBox_30.value(),
            F_ax=self.doubleSpinBox_71.value(),
            delta_phi=self.doubleSpinBox_28.value(),
            delta_alpha=self.doubleSpinBox_29.value(),
            alpha=self.doubleSpinBox_45.value(),
            W_dp=self.doubleSpinBox_49.value(),
            L_dp=self.doubleSpinBox_50.value()
        )

        self.d_m = wear_model_line.d
        print(f"d_m = {self.d_m*1000} mm\n")

        length_of_signal = 50
        t = np.linspace(5, 50, length_of_signal)
        FS_signal = []
        for i in range(length_of_signal):
            FS = cm.Wear_model_line(
                mu=self.doubleSpinBox_26.value(),
                n=self.doubleSpinBox_27.value(),
                f_w=self.doubleSpinBox_25.value(),
                D=self.doubleSpinBox_24.value(),
                L_m=self.doubleSpinBox_22.value(),
                Rc=self.doubleSpinBox_30.value(),
                F_ax=self.doubleSpinBox_71.value(),
                delta_phi=self.doubleSpinBox_28.value(),
                delta_alpha=self.doubleSpinBox_29.value(),
                alpha=self.doubleSpinBox_45.value(),
                W_dp=self.doubleSpinBox_49.value(),
                L_dp=self.doubleSpinBox_50.value(),
                v_rop=t[i]
            )
            FS_signal.append(FS.S)

        print(FS_signal)

        # 设定字体为微软雅黑
        plt.rcParams['font.sans-serif'] = ['Microsoft Yahei']

        # 清除当前图形以便绘制新的信号
        self.widget_2.canvas.axes.clear()
        # 绘制信号
        self.widget_2.canvas.axes.plot(t, FS_signal)
        # 设置x轴标签
        self.widget_2.canvas.axes.set_xlabel('钻速（m/h）')  # 设置图像标签
        # 设置y轴标签
        self.widget_2.canvas.axes.set_ylabel('磨损面积（m^2）')
        # 设置图形标题
        self.widget_2.canvas.axes.set_title(' 磨损模型 ')

        # 绘制图形
        self.widget_2.canvas.draw()

        # 保存为图片文件
        self.widget_2.canvas.figure.savefig(self.temp_image_path[1], dpi=300, bbox_inches='tight')

    def corrode_model_graph(self):  # 更新图像数据
        """
        更新图形界面中的数据。

        """
        #  腐蚀模型
        corrode_model = cm.corrode_model(
            T=self.doubleSpinBox.value(),
            P_co2=self.doubleSpinBox_32.value(),
            P_h2s=self.doubleSpinBox_34.value(),
            Cl=self.doubleSpinBox_33.value(),
            pH=self.doubleSpinBox_35.value(),
            material=self.comboBox_2.currentText(),
            t=self.doubleSpinBox_54.value()
        )

        # 计算腐蚀壁厚
        self.d_f = corrode_model.d_f
        print(f"d_f = {self.d_f} mm\n")

        length_of_signal = 50
        t = np.linspace(0.001, 100, length_of_signal)
        FS_signal = []
        if self.doubleSpinBox_34.value() == 0:
            for i in range(length_of_signal):
                FS = cm.corrode_model(
                    T=t[i],
                    P_co2=self.doubleSpinBox_32.value(),
                    P_h2s=self.doubleSpinBox_34.value(),
                    Cl=self.doubleSpinBox_33.value(),
                    pH=self.doubleSpinBox_35.value(),
                    material=self.comboBox_2.currentText(),
                    t=self.doubleSpinBox_54.value()
                )
                FS_signal.append(FS.R_year)

            print(FS_signal)

            # 设定字体为微软雅黑
            plt.rcParams['font.sans-serif'] = ['Microsoft Yahei']

            # 清除当前图形以便绘制新的信号
            self.graphicsView_3.canvas.axes.clear()
            # 绘制信号
            self.graphicsView_3.canvas.axes.plot(t, FS_signal)
            # 设置x轴标签
            self.graphicsView_3.canvas.axes.set_xlabel('温度（℃）')  # 设置图像标签
            # 设置y轴标签
            self.graphicsView_3.canvas.axes.set_ylabel('腐蚀速率（mm/year）')
            # 设置图形标题
            self.graphicsView_3.canvas.axes.set_title('二氧化碳环境腐蚀预测模型 ')

            # 绘制图形
            self.graphicsView_3.canvas.draw()
        elif self.doubleSpinBox_34.value() != 0:
            for i in range(length_of_signal):
                FS = cm.corrode_model(
                    T=t[i],
                    P_co2=self.doubleSpinBox_32.value(),
                    P_h2s=self.doubleSpinBox_34.value(),
                    Cl=self.doubleSpinBox_33.value(),
                    pH=self.doubleSpinBox_35.value(),
                    material=self.comboBox_2.currentText(),
                    t=self.doubleSpinBox_54.value()
                )
                FS_signal.append(FS.R_year)

            print(FS_signal)

            # 设定字体为微软雅黑
            plt.rcParams['font.sans-serif'] = ['Microsoft Yahei']

            # 清除当前图形以便绘制新的信号
            self.graphicsView_3.canvas.axes.clear()
            # 绘制信号
            self.graphicsView_3.canvas.axes.plot(t, FS_signal)
            # 设置x轴标签
            self.graphicsView_3.canvas.axes.set_xlabel('温度（℃）')  # 设置图像标签
            # 设置y轴标签
            self.graphicsView_3.canvas.axes.set_ylabel('腐蚀速率（mm/year）')
            # 设置图形标题
            self.graphicsView_3.canvas.axes.set_title('二氧化碳与硫化氢共存环境腐蚀预测模型 ')

            # 绘制图形
            self.graphicsView_3.canvas.draw()

            # 保存为图片文件
            self.graphicsView_3.canvas.figure.savefig(self.temp_image_path[2], dpi=300, bbox_inches='tight')

            # 显示长期腐蚀速率
            self.lineEdit_8.setText(f"{corrode_model.R_year * 10000:.4f}")

            self.calc_result['长期腐蚀速率'] = corrode_model.R_year * 10000

    def noplasticity_effective_external_squeeze_pressure_graph(self):  # 更新图像数据
        """
        更新图形界面中的数据。

        """
        # 失效判定——非蠕变地层最大外压力
        self.no_plasticity_effective_external_squeeze_pressure = cm.NoPlasticity_Effective_external_squeeze_pressure(
            rho_m=self.doubleSpinBox_56.value(),
            rho_w=self.doubleSpinBox_65.value(),
            k_m=self.doubleSpinBox_57.value(),
            rho_min=self.doubleSpinBox_59.value(),
            casing_type=self.comboBox_5.currentText(),
            h=self.doubleSpinBox_58.value()
        )
        self.P_ce = self.no_plasticity_effective_external_squeeze_pressure.p_ce

        length_of_signal = 50
        t = np.linspace(1, 10000, length_of_signal)
        FS_signal = []
        for i in range(length_of_signal):
            FS = cm.NoPlasticity_Effective_external_squeeze_pressure(
                rho_m=self.doubleSpinBox_56.value(),
                rho_w=self.doubleSpinBox_65.value(),
                k_m=self.doubleSpinBox_57.value(),
                rho_min=self.doubleSpinBox_59.value(),
                casing_type=self.comboBox_5.currentText(),
                h=t[i]
            )
            FS_signal.append(FS.p_ce)

        print(FS_signal)

        # 设定字体为微软雅黑
        plt.rcParams['font.sans-serif'] = ['Microsoft Yahei']

        # 清除当前图形以便绘制新的信号
        self.widget_6.canvas.axes.clear()
        # 绘制信号
        self.widget_6.canvas.axes.plot(t, FS_signal)
        # 设置x轴标签
        self.widget_6.canvas.axes.set_xlabel('计算点深度（m）')  # 设置图像标签
        # 设置y轴标签
        self.widget_6.canvas.axes.set_ylabel('最大外压力（Mpa）')
        # 设置图形标题
        self.widget_6.canvas.axes.set_title(' 最大外压力模型 ')

        # 绘制图形
        self.widget_6.canvas.draw()

        # 保存为图片文件
        self.widget_6.canvas.figure.savefig(self.temp_image_path[3], dpi=300, bbox_inches='tight')

        # 显示最大外压力
        self.lineEdit_6.setText(f"{self.no_plasticity_effective_external_squeeze_pressure.p_ce:.2f}")

        self.calc_result['最大外压力'] = self.no_plasticity_effective_external_squeeze_pressure.p_ce

    def plasticity_effective_external_squeeze_pressure_graph(self):  # 更新图像数据
        """
        更新图形界面中的数据。

        """
        # 失效判定——蠕变地层最大外压力
        self.plasticity_effective_external_squeeze_pressure = cm.Plasticity_Effective_external_squeeze_pressure(
            k_m=self.doubleSpinBox_62.value(),
            rho_min=self.doubleSpinBox_60.value(),
            rho_w=self.doubleSpinBox_66.value(),
            h=self.doubleSpinBox_61.value(),
            v=self.doubleSpinBox_63.value(),
            casing_type=self.comboBox_6.currentText(),
            G_v=0.023
        )
        self.P_ce = self.plasticity_effective_external_squeeze_pressure.p_ce

        length_of_signal = 50
        t = np.linspace(1, 10000, length_of_signal)
        FS_signal = []
        for i in range(length_of_signal):
            FS = cm.Plasticity_Effective_external_squeeze_pressure(
                k_m=self.doubleSpinBox_62.value(),
                rho_min=self.doubleSpinBox_60.value(),
                rho_w=self.doubleSpinBox_66.value(),
                h=t[i],
                v=self.doubleSpinBox_63.value(),
                casing_type=self.comboBox_6.currentText(),
                G_v=0.023
            )
            FS_signal.append(FS.p_ce)

        print(FS_signal)

        # 设定字体为微软雅黑
        plt.rcParams['font.sans-serif'] = ['Microsoft Yahei']

        # 清除当前图形以便绘制新的信号
        self.widget_7.canvas.axes.clear()
        # 绘制信号
        self.widget_7.canvas.axes.plot(t, FS_signal)
        # 设置x轴标签
        self.widget_7.canvas.axes.set_xlabel('计算点深度（m）')  # 设置图像标签
        # 设置y轴标签
        self.widget_7.canvas.axes.set_ylabel('最大外压力（Mpa）')
        # 设置图形标题
        self.widget_7.canvas.axes.set_title(' 最大外压力模型 ')

        # 绘制图形
        self.widget_7.canvas.draw()

        # 保存为图片文件
        self.widget_7.canvas.figure.savefig(self.temp_image_path[3], dpi=300, bbox_inches='tight')

        # 显示最大外压力
        self.lineEdit_7.setText(f"{self.plasticity_effective_external_squeeze_pressure.p_ce:.2f}")

        self.calc_result['最大外压力'] = self.plasticity_effective_external_squeeze_pressure.p_ce

    def gas_effective_internal_pressure_graph(self):  # 更新图像数据
        """
        更新图形界面中的数据。

        """
        # 失效判定——气井最大内压力
        self.gas_Effective_internal_pressure = cm.Gas_Effective_internal_pressure(
            rho_max=self.doubleSpinBox_36.value(),
            H_s=self.doubleSpinBox_37.value(),
            p_p=self.doubleSpinBox_38.value(),
            rho_g=self.doubleSpinBox_40.value(),
            h=self.doubleSpinBox_39.value(),
            H_mg=self.doubleSpinBox_41.value(),
            casing_type=self.comboBox_3.currentText()
        )
        self.P_bh = self.gas_Effective_internal_pressure.p_bh

        length_of_signal = 50
        t = np.linspace(1, 10000, length_of_signal)
        FS_signal = []
        for i in range(length_of_signal):
            FS = cm.Gas_Effective_internal_pressure(
                rho_max=self.doubleSpinBox_36.value(),
                H_s=self.doubleSpinBox_37.value(),
                p_p=self.doubleSpinBox_38.value(),
                rho_g=self.doubleSpinBox_40.value(),
                H_mg=self.doubleSpinBox_41.value(),
                casing_type=self.comboBox_3.currentText(),
                h=t[i]
            )
            FS_signal.append(FS.p_bh)

        print(FS_signal)

        # 设定字体为微软雅黑
        plt.rcParams['font.sans-serif'] = ['Microsoft Yahei']

        # 清除当前图形以便绘制新的信号
        self.widget_8.canvas.axes.clear()
        # 绘制信号
        self.widget_8.canvas.axes.plot(t, FS_signal)
        # 设置x轴标签
        self.widget_8.canvas.axes.set_xlabel('计算点深度（m）')  # 设置图像标签
        # 设置y轴标签
        self.widget_8.canvas.axes.set_ylabel('最大内挤压力（Mpa）')
        # 设置图形标题
        self.widget_8.canvas.axes.set_title(' 最大内挤压力模型 ')

        # 绘制图形
        self.widget_8.canvas.draw()

        # 保存为图片文件
        self.widget_8.canvas.figure.savefig(self.temp_image_path[4], dpi=300, bbox_inches='tight')

        # 显示最大内压力
        self.lineEdit_4.setText(f"{self.gas_Effective_internal_pressure.p_bh:.2f}")

        self.calc_result['最大内压力'] = self.gas_Effective_internal_pressure.p_bh

    def oil_effective_internal_squeeze_pressure_graph(self):  # 更新图像数据
        """
        更新图形界面中的数据。

        """
        # 失效判定——油井最大内压力
        self.oil_Effective_internal_squeeze_pressure = cm.Oil_Effective_internal_squeeze_pressure(
            rho_max=self.doubleSpinBox_42.value(),
            rho_w=self.doubleSpinBox_46.value(),
            G=self.doubleSpinBox_47.value(),
            h=self.doubleSpinBox_43.value(),
            H_s=self.doubleSpinBox_44.value(),
            casing_type=self.comboBox_4.currentText()
        )
        self.P_bh = self.oil_Effective_internal_squeeze_pressure.p_be

        length_of_signal = 50
        t = np.linspace(0.1, 10000, length_of_signal)
        FS_signal = []
        for i in range(length_of_signal):
            FS = cm.Oil_Effective_internal_squeeze_pressure(
                rho_max=self.doubleSpinBox_42.value(),
                rho_w=self.doubleSpinBox_46.value(),
                G=self.doubleSpinBox_47.value(),
                h=t[i],
                H_s=self.doubleSpinBox_44.value(),
                casing_type=self.comboBox_4.currentText()
            )
            FS_signal.append(FS.p_be)

        print(FS_signal)

        # 设定字体为微软雅黑
        plt.rcParams['font.sans-serif'] = ['Microsoft Yahei']

        # 清除当前图形以便绘制新的信号
        self.widget_9.canvas.axes.clear()
        # 绘制信号
        self.widget_9.canvas.axes.plot(t, FS_signal)
        # 设置x轴标签
        self.widget_9.canvas.axes.set_xlabel('计算点深度（m）')  # 设置图像标签
        # 设置y轴标签
        self.widget_9.canvas.axes.set_ylabel('最大内挤压力（Mpa）')
        # 设置图形标题
        self.widget_9.canvas.axes.set_title(' 最大内挤压力模型 ')

        # 绘制图形
        self.widget_9.canvas.draw()

        # 保存为图片文件
        self.widget_9.canvas.figure.savefig(self.temp_image_path[4], dpi=300, bbox_inches='tight')

        # 显示最大内挤压力
        self.lineEdit_5.setText(f"{self.oil_Effective_internal_squeeze_pressure.p_be:.2f}")

        self.calc_result['最大内压力'] = self.oil_Effective_internal_squeeze_pressure.p_be

    def effective_internal_pressure_failure_condition_calculate(self):  # 更新图像数据
        """
        更新图形界面中的数据。

        """
        try:
            if not hasattr(self, 'P_ce') or not hasattr(self, 'P_bh'):
                QMessageBox.warning(self, "提示", "请先计算")
                return

            casing = self.comboBox_7.currentText()
            if casing == "P110":
                Y_p = 800
            else:
                Y_p = 600
            # 套管在内压、外压作用下失效判定条件
            self.effective_internal_pressure_failure_condition = cm.Effective_internal_pressure_failure_condition(
                d=self.doubleSpinBox_55.value(),
                d_c=self.d_c,
                d_m=self.d_m,
                d_f=self.d_f,
                Y_p=Y_p,
                D=self.doubleSpinBox_52.value()
            )

            # 显示剩余抗内挤强度
            self.lineEdit_3.setText(f"{self.effective_internal_pressure_failure_condition.P_bo:.2f}")

            self.calc_result['剩余抗内挤强度'] = self.effective_internal_pressure_failure_condition.P_bo

            # 显示剩余抗外挤强度
            self.lineEdit_2.setText(f"{self.effective_internal_pressure_failure_condition.P_co:.2f}")

            self.calc_result['剩余抗外挤强度'] = self.effective_internal_pressure_failure_condition.P_co

            if self.P_ce > 0 and self.P_bh > 0:
                S_W = self.effective_internal_pressure_failure_condition.P_co / self.P_ce
                S_N = self.effective_internal_pressure_failure_condition.P_bo / self.P_bh

                # 显示安全等级
                if (S_W > 1.125 and S_N > 1.25):
                    self.lineEdit.setText("安全")
                    self.lineEdit.setStyleSheet("color: green;")
                    self.calc_result['安全等级'] = "安全"
                elif (1 < S_W < 1.125 or 1 < S_N < 1.25):
                    self.lineEdit.setText("警告")
                    self.lineEdit.setStyleSheet("color: Gold;")
                    self.calc_result['安全等级'] = "警告"
                else:
                    self.lineEdit.setText("危险")
                    self.lineEdit.setStyleSheet("color: red;")
                    self.calc_result['安全等级'] = "危险"
            else:
                QMessageBox.warning(self, "提示", "请先计算")
                raise ValueError("剩余强度必须大于0")

        except Exception as e:
            print(f"求解过程中发生错误: {str(e)}")

    def open_excel_file(self):
        """
            打开 Excel 文件并将数据加载对应控件中

        """
        # 打开文件对话框选择 Excel 文件
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择 Excel 文件", "输入数据.xlsx", "Excel Files (*.xlsx *.xls);;All Files (*)"
        )

        if not file_path:
            return  # 用户取消选择

        try:
            # 使用 pandas 读取 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为空
            if df.empty:
                QMessageBox.warning(self, "警告", "Excel 文件中没有数据!")
                return

            dialog = ExcelDataSelector()
            rows = df.iloc[0:50]
            print(rows)

            if dialog.exec_():
                selections = dialog.getSelections()
                # 获取用户选择
                well_type = selections["well_type"]
                wear_model = selections["wear_model"]
                # environment = selections["environment"]
                formation = selections["formation"]
                # 根据选择加载对应的数据
                if well_type == "气井数据":
                    # 从Excel中加载气井数据
                    '''填充冲蚀数据'''
                    self.doubleSpinBox_3.setValue(rows.iloc[0, 1])
                    self.doubleSpinBox_2.setValue(rows.iloc[1, 1])
                    self.doubleSpinBox_5.setValue(rows.iloc[2, 1])

                    if rows.iloc[3, 1] == "圆形":
                        self.comboBox.setCurrentText("圆形")
                    elif rows.iloc[3, 1] == "半圆形":
                        self.comboBox.setCurrentText("半圆形")
                    else:
                        self.comboBox.setCurrentText("角形")

                    self.doubleSpinBox_67.setValue(rows.iloc[4, 1])
                    self.doubleSpinBox_68.setValue(rows.iloc[5, 1])
                    self.doubleSpinBox_69.setValue(rows.iloc[6, 1])
                    '''填充最大内压力数据'''
                    if rows.iloc[25, 1] == "表层套管和技术套管":
                        self.comboBox_3.setCurrentText("表层套管和技术套管")
                    else:
                        self.comboBox_3.setCurrentText("生产套管和生产尾管")

                    self.doubleSpinBox_36.setValue(rows.iloc[32, 1])
                    self.doubleSpinBox_40.setValue(rows.iloc[33, 1])
                    self.doubleSpinBox_39.setValue(rows.iloc[30, 1])
                    self.doubleSpinBox_41.setValue(rows.iloc[40, 1])
                    self.doubleSpinBox_37.setValue(rows.iloc[34, 1])
                    self.doubleSpinBox_38.setValue(rows.iloc[35, 1])

                else:  # 油井数据
                    '''填充冲蚀数据'''
                    self.doubleSpinBox_8.setValue(rows.iloc[1, 1])
                    self.doubleSpinBox_64.setValue(rows.iloc[7, 1])
                    self.doubleSpinBox_12.setValue(rows.iloc[6, 1])
                    self.doubleSpinBox_13.setValue(rows.iloc[2, 1])
                    self.doubleSpinBox_20.setValue(rows.iloc[8, 1])
                    self.doubleSpinBox_19.setValue(rows.iloc[9, 1])
                    self.doubleSpinBox_17.setValue(rows.iloc[5, 1])
                    self.doubleSpinBox_21.setValue(rows.iloc[10, 1])
                    '''填充最大内压力'''
                    if rows.iloc[25, 1] == "表层套管和技术套管":
                        self.comboBox_4.setCurrentText("表层套管和技术套管")
                    else:
                        self.comboBox_4.setCurrentText("生产套管和生产尾管")

                    self.doubleSpinBox_42.setValue(rows.iloc[32, 1])
                    self.doubleSpinBox_46.setValue(rows.iloc[29, 1])
                    self.doubleSpinBox_47.setValue(rows.iloc[36, 1])
                    self.doubleSpinBox_43.setValue(rows.iloc[30, 1])
                    self.doubleSpinBox_44.setValue(rows.iloc[34, 1])

                # 根据磨损模型加载不同的参数
                if wear_model == "套管磨损面积(线性关系)":
                    self.doubleSpinBox_26.setValue(rows.iloc[11, 1])
                    self.doubleSpinBox_27.setValue(rows.iloc[12, 1])
                    self.doubleSpinBox_25.setValue(rows.iloc[13, 1])
                    self.doubleSpinBox_24.setValue(rows.iloc[14, 1])
                    self.doubleSpinBox_22.setValue(rows.iloc[15, 1])
                    self.doubleSpinBox_23.setValue(rows.iloc[16, 1])
                    self.doubleSpinBox_30.setValue(rows.iloc[39, 1]/2)
                    self.doubleSpinBox_71.setValue(rows.iloc[41, 1])
                    self.doubleSpinBox_28.setValue(rows.iloc[42, 1])
                    self.doubleSpinBox_29.setValue(rows.iloc[43, 1])
                    self.doubleSpinBox_45.setValue(rows.iloc[44, 1])
                    self.doubleSpinBox_49.setValue(rows.iloc[45, 1])
                    self.doubleSpinBox_50.setValue(rows.iloc[46, 1])

                else:  # 几何关系
                    self.doubleSpinBox_28.setValue(rows.iloc[14, 1] / 2)
                    self.doubleSpinBox_30.setValue(rows.iloc[17, 1])

                # 为腐蚀环境加载不同的参数
                if rows.iloc[18, 1] == "碳钢":
                    self.comboBox.setCurrentText("碳钢")
                elif rows.iloc[18, 1] == "1Cr":
                    self.comboBox.setCurrentText("1Cr")
                else:
                    self.comboBox.setCurrentText("3Cr")

                self.doubleSpinBox.setValue(rows.iloc[19, 1])
                self.doubleSpinBox_32.setValue(rows.iloc[20, 1])
                self.doubleSpinBox_34.setValue(rows.iloc[21, 1])
                self.doubleSpinBox_33.setValue(rows.iloc[22, 1])
                self.doubleSpinBox_35.setValue(rows.iloc[23, 1])
                self.doubleSpinBox_54.setValue(rows.iloc[24, 1])

                # 根据地层类型加载不同的参数
                if formation == "非塑性蠕变地层":
                    if rows.iloc[25, 1] == "表层套管和技术套管":
                        self.comboBox_5.setCurrentText("表层套管和技术套管")
                    else:
                        self.comboBox_5.setCurrentText("生产套管和生产尾管")

                    self.doubleSpinBox_56.setValue(rows.iloc[24, 1])
                    self.doubleSpinBox_57.setValue(rows.iloc[27, 1])
                    self.doubleSpinBox_59.setValue(rows.iloc[28, 1])
                    self.doubleSpinBox_65.setValue(rows.iloc[29, 1])
                    self.doubleSpinBox_58.setValue(rows.iloc[30, 1])
                else:  # 塑性蠕变地层
                    if rows.iloc[25, 1] == "表层套管和技术套管":
                        self.comboBox_6.setCurrentText("表层套管和技术套管")
                    else:
                        self.comboBox_6.setCurrentText("生产套管和生产尾管")

                    self.doubleSpinBox_60.setValue(rows.iloc[28, 1])
                    self.doubleSpinBox_66.setValue(rows.iloc[29, 1])
                    self.doubleSpinBox_62.setValue(rows.iloc[27, 1])
                    self.doubleSpinBox_63.setValue(rows.iloc[31, 1])
                    self.doubleSpinBox_61.setValue(rows.iloc[30, 1])

                self.doubleSpinBox_48.setValue(rows.iloc[6, 1])
                self.doubleSpinBox_51.setValue(rows.iloc[24, 1])
                if rows.iloc[37, 1] == "N80":
                    self.comboBox_7.setCurrentText("N80")
                else:
                    self.comboBox_7.setCurrentText("P110")

                self.doubleSpinBox_52.setValue(rows.iloc[39, 1])
                self.doubleSpinBox_55.setValue(rows.iloc[38, 1])

                # 显示成功消息
                QMessageBox.information(self, "成功", f"已成功加载 Excel 文件")
                self.statusBar().showMessage("已将选中数据导入到参数输入框")

                for key, value in selections.items():
                    print(f"{key}: {value}")
            else:
                print("用户取消了操作")

        except Exception as e:
            # 错误处理
            QMessageBox.critical(self, "错误", f"打开文件时出错: {str(e)}")

    def export_to_template(self):
        """使用模板导出结果"""
        if any(v is None for v in self.calc_result.values()):
            QMessageBox.warning(self, "提示", "请先计算")
            return

        # # 保存当前图表
        # self.figure.savefig(self.temp_image_path)

        # 选择模板文件
        template_path, _ = QFileDialog.getOpenFileName(
            self, "选择Word模板", "", "Word文档 (*.docx)"
        )

        if not template_path:
            return

        # 选择输出路径
        output_path, _ = QFileDialog.getSaveFileName(
            self, "保存结果文档", "", "Word文档 (*.docx)"
        )

        if not output_path:
            return

        try:
            # 使用模板填充结果
            success, message, output_file = fill_template_with_results(
                template_path,
                self.calc_result,
                self.temp_image_path,
                output_path
            )

            if success:
                self.last_template_output = output_file  # 记录模板导出的文件路径
                QMessageBox.information(self, "成功", f"已使用模板生成文档:\n{output_path}")
            else:
                QMessageBox.warning(self, "警告", message)

        except Exception as e:
            QMessageBox.critical(self, "错误", f"生成文档失败: {str(e)}")

    def closeEvent(self, event):
        """清理临时文件"""
        for file_path in self.temp_image_path:
            # 确保路径是字符串且文件存在
            if isinstance(file_path, str) and os.path.exists(file_path):
                os.remove(file_path)
        event.accept()

    def calculate_all(self):

        return
